# -*- coding: utf-8 -*-
"""
Created on Tue Sep 19 14:10:44 2017

@author: choip
"""
import docx

doc = docx.Document('demo.docx')
#print (len(doc.paragraphs))
#7 paragraphs in doc
#print (doc.paragraphs[0].text)
#print (doc.paragraphs[1].text)
#print (doc.paragraphs[1].runs[0].text)
#print (doc.paragraphs[1].runs[1].text)
#print (doc.paragraphs[1].runs[2].text)
#print (doc.paragraphs[1].runs[3].text)
#run is a list of words that are split up 
#according to style
#for example, above says a plain paragraph with some
#bold, because bold is boldened
#and some
#italic, because italic is italicized. 

#define a function to get all text in a doc

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

#to indent each pargraph
#fullText.append(' ' + para.text)
#to add a double space in between paragraphs
#return '\n\n'.join(fullText)

print (getText('demo.docx'))

#Run, characters, attributes
#Attribute
#
#Description
#
#bold
#
#The text appears in bold.
#
#italic
#
#The text appears in italic.
#
#underline
#
#The text is underlined.
#
#strike
#
#The text appears with strikethrough.
#
#double_strike
#
#The text appears with double strikethrough.
#
#all_caps
#
#The text appears in capital letters.
#
#small_caps
#
#The text appears in capital letters, with lowercase letters two points smaller.
#
#shadow
#
#The text appears with a shadow.
#
#outline
#
#The text appears outlined rather than solid.
#
#rtl
#
#The text is written right-to-left.
#
#imprint
#
#The text appears pressed into the page.
#
#emboss
#
#The text appears raised off the page in relief.

#doc.paragraphs[0].style = 'Normal'
#doc.paragraphs[1].runs[0].style = 'QuoteChar'
##end with char if you want to modify runs
#doc.paragraphs[1].runs[1].underline = True
#doc.paragraphs[1].runs[3].underline = True
#doc.save('restyled.docx')

#doc = docx.Document()
#doc.add_paragraph('Hello world!')
#doc.save('helloworld.docx')
#paraObj1 = doc.add_paragraph('This is a second paragraph.')
#paraObj2 = doc.add_paragraph('This is another paragraph.')
#paraObj1.add_run(' This text is being added to the second paragraph')
#doc.save('multipleParagraphs.docx')

#doc = docx.Document()
#doc.add_heading('Header 0',0)
#doc.add_heading('Header2',2)
#doc.add_heading('Header4',4)
#doc.save('heading.docx')

#doc = docx.Document()
#doc.add_paragraph('This is on the first page!')
#doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
#doc.save('twoPage.docx')
#don't forget enum.

#doc.add_picture('zophie.png', width=docx.shared.Inches(1),
#height=docx.shared.Cm(4))
#docx can also take centimer docx.shared.Cm()